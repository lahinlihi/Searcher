"""
LH API 문서 추출 (테이블 포함)
"""

from docx import Document

doc = Document('D:/API3.docx')

output = []

# 본문 텍스트
for para in doc.paragraphs:
    if para.text.strip():
        output.append(para.text)

# 테이블
for i, table in enumerate(doc.tables):
    output.append(f"\n[테이블 {i+1}]")
    for row in table.rows:
        row_text = ' | '.join([cell.text.strip() for cell in row.cells])
        output.append(row_text)

# 파일로 저장
with open('lh_api_doc_full.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print(f"총 {len(output)}줄 추출 완료")
print("파일 저장: lh_api_doc_full.txt")
